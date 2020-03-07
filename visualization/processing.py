from docx import Document
import re
from docx.shared import Pt
from docx.shared import RGBColor
import json
import pandas as pd
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_COLOR_INDEX, WD_BREAK
from nltk.corpus import stopwords
import re
import math
from docx.shared import Cm, Inches
import csv

stop_words = set(stopwords.words('english'))
df_res = pd.read_csv('results_temp7.csv')
with open('cluster_properties_temp7.json', 'r') as f:
    cluster_props = json.load(f)

file_class_one = 'class_one3.docx'
file_class_two = 'class_two3.docx'
file_class_three = 'class_three3.docx'

csv_class1 = 'class_one.csv'
csv_class2 = 'class_two.csv'
csv_class3 = 'class_three.csv'

class1_clusters = []
class2_clusters = []
class3_clusters = []

clusters = df_res['cluster_label'].unique()

cluster_word_list = []
avg_label_scores = []
cluster_metric = []
top_bigrams = []

for cl in clusters:
    if cl == -1:
        avg_label_scores.append(0)
        cluster_metric.append(0)
        cluster_word_list.append('')
        top_bigrams.append('')
        continue
    cl = str(cl)
    avg_label_scores.append(cluster_props[cl]['avg_label_score'])
    cluster_metric.append(cluster_props[cl]['cluster_metric'])
    cluster_word_list.append(cluster_props[cl]['word_list'])
    top_bigrams.append(cluster_props[cl]['top_bigrams'])
    # top_bigrams.append('')

cluster_metric, avg_label_scores, clusters, cluster_word_list, top_bigrams = zip(*sorted(zip(cluster_metric, avg_label_scores, clusters, 
    cluster_word_list, top_bigrams)))

cluster_word_list = list(reversed(cluster_word_list))
clusters = list(reversed(clusters))
avg_label_scores = list(reversed(avg_label_scores))
cluster_metric = list(reversed(cluster_metric))
top_bigrams = list(reversed(top_bigrams))

doc1 = Document()
doc2 = Document()
doc3 = Document()

doc1.add_heading('Class 1 clusters')
doc2.add_heading('Class 2 clusters')
doc3.add_heading('Class 3 clusters')

section = doc1.sections[0]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

section = doc2.sections[0]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

section = doc3.sections[0]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height
# 2 - blue, 4 - bright green, 5 - dark_blue, 6 - dark red, 7 - dark yellow
# 8 - gray, 10 - green, 11 -pink, 12 - red, 13- teal, 14- turquoise, 15- violet

#141414 - black, FF0000 - red, F3FF69 - yellow, FFACB7 - pink, 808000- mehendi, 6897BB - blue

FONT_MAP = {
    0: {'size' : Pt(10), 'highlight' : WD_COLOR_INDEX.BRIGHT_GREEN, 'color': RGBColor.from_string('141414')},
    1: {'size' : Pt(10), 'highlight' : WD_COLOR_INDEX.PINK, 'color': RGBColor.from_string('141414')},
    2: {'size' : Pt(10), 'highlight' : WD_COLOR_INDEX.RED, 'color': RGBColor.from_string('141414')},
    3: {'size' : Pt(10), 'highlight' : WD_COLOR_INDEX.TURQUOISE, 'color': RGBColor.from_string('141414')},
    4: {'size' : Pt(10), 'highlight' : WD_COLOR_INDEX.YELLOW, 'color': RGBColor.from_string('141414')},
    5: {'size' : Pt(10), 'highlight' : WD_COLOR_INDEX.GRAY_25, 'color': RGBColor.from_string('141414')},
    6: {'size' : Pt(9), 'highlight' : WD_COLOR_INDEX.BRIGHT_GREEN, 'color': RGBColor.from_string('6897BB')},
    7: {'size' : Pt(9), 'highlight' : WD_COLOR_INDEX.PINK, 'color': RGBColor.from_string('6897BB')},
    8: {'size' : Pt(9), 'highlight' : WD_COLOR_INDEX.RED, 'color': RGBColor.from_string('FFACB7')},
    9: {'size' : Pt(9), 'highlight' : WD_COLOR_INDEX.TURQUOISE, 'color': RGBColor.from_string('6897BB')},
    10: {'size' : Pt(9), 'highlight' : WD_COLOR_INDEX.YELLOW, 'color': RGBColor.from_string('6897BB')},
    11: {'size' : Pt(8), 'highlight' : WD_COLOR_INDEX.GRAY_25, 'color': RGBColor.from_string('6897BB')},
    12: {'size' : Pt(8), 'highlight' : WD_COLOR_INDEX.BRIGHT_GREEN, 'color': RGBColor.from_string('FF0000')},
    13: {'size' : Pt(8), 'highlight' : WD_COLOR_INDEX.PINK, 'color': RGBColor.from_string('FF0000')},
    14: {'size' : Pt(8), 'highlight' : WD_COLOR_INDEX.RED, 'color': RGBColor.from_string('F3FF69')},
    15: {'size' : Pt(8), 'highlight' : WD_COLOR_INDEX.TURQUOISE, 'color': RGBColor.from_string('FF0000')},
    16: {'size' : Pt(7), 'highlight' : WD_COLOR_INDEX.YELLOW, 'color': RGBColor.from_string('FF0000')},
    17: {'size' : Pt(7), 'highlight' : WD_COLOR_INDEX.GRAY_25, 'color': RGBColor.from_string('FF0000')},
    18: {'size' : Pt(7), 'highlight' : WD_COLOR_INDEX.BRIGHT_GREEN, 'color': RGBColor.from_string('FFACB7')},
    19: {'size' : Pt(7), 'highlight' : WD_COLOR_INDEX.PINK, 'color': RGBColor.from_string('F3FF69')},
    20: {'size' : Pt(7), 'highlight' : WD_COLOR_INDEX.RED, 'color': RGBColor.from_string('F3FF69')},
}

def get_start_tag(count):
    return "HSTART_{}".format(count)

def get_end_tag(count):
    return "HEND"

def get_highlighted_body(body, word_list):
    count = 0
    for i, bigram in enumerate(word_list):
#             print (bigram)
        try:
            start_ind = body.lower().index(bigram)
            end_ind = start_ind + len(bigram)
#                 print (start_ind)
            body = body[:start_ind] + get_start_tag(i%10) + body[start_ind:end_ind] + get_end_tag(i%10) + body[end_ind:]
            count += 1
            if count ==10:
                break
        except:
            pass

    return body

# def unify_phrases(body):
#     new_body = ''
#     # r = re.compile('HSTART_')
#     # fragments = r.split(h_body)
#     # new_body = fragments[0]
#     count = 0
#     new_count = 0
#     split_b = body.split(get_start_tag(count))
#     new_body = split_b[0] + get_start_tag(new_count) + ' '
#     sub_split = split_b[1].split(get_end_tag(count))
#     first_part = sub_split[0]
#     new_body += first_part
#     while True:
#         count += 1
#         if get_start_tag(count) in sub_split[1]:
#             next_phrase = sub_split[1].split(get_start_tag(count))
#             second_part = next_phrase[0]
#             sub_split = split_b[1].split(get_end_tag(count))
#             third_part = sub_split[0]
#             should_merge = check_merging_criteria(second)
#             if should_merge:
#                 new_body += second_part + third_part
#             else:
#                 new_body += (' ' + get_end_tag(new_count) + second_part)
#                 new_count += 1
#                 new_body += (get_start_tag(new_count)+ ' ')
#         else:
#             new_body += (' ' + get_end_tag(new_count) + sub_split[1])
#             break

#     return new_body

def should_merge(second):
    second = re.sub('[^0-9a-zA-Z ]+', '', second)
    second = second.strip()
    if second is None or second == '' or second == ' ':
        return True
    all_words = second.split()
    stopword_count = 0
    normal_word_count = 0

    for w in all_words:
        if len(w) < 4:
            continue
        if w in stop_words:
            stopword_count += 1
        else:
            normal_word_count += 1

    if normal_word_count == 0:
        return True

    return False

def unify_phrases(body):
    # print (body)
    r = re.compile('HSTART_')
    fragments = r.split(body)
    # print(body)
    if len(fragments) == 1:
        return body
    new_body = fragments[0]
    new_count = 0
    color_index = int(fragments[1].split('HEND')[0][0])
    first = fragments[1].split('HEND')[0][1:]
    second = fragments[1].split('HEND')[1]
    new_body += (get_start_tag(color_index)) + first
    ended = False
    # print (len(fragments[2:]))
    for index, frag in enumerate(fragments[2:]):
        # print (frag)
        # if index == len(fragments[2:])-1:
        #     print ('=-------------------------------')
        #     print (first, second)
        #     new_body += second + get_end_tag(new_count)
        #     first = frag.split('HEND')[0][1:]
        #     second = frag.split('HEND')[1]
        if should_merge(second):
            # ended = False
            new_body += second
            first = frag.split('HEND')[0][1:]
            second = frag.split('HEND')[1]
            # print (frag)
            # print(first)
            # exit()
            new_body += first
        else:
            # ended = True
            new_body += get_end_tag(color_index) + second
            new_count += 1
            color_index = int(frag.split('HEND')[0][0])
            new_body += get_start_tag(color_index)
            first = frag.split('HEND')[0][1:]
            second = frag.split('HEND')[1]
            new_body += first

    new_body += get_end_tag(color_index) + second
    #         new_body += first

    return new_body

def get_para_object(count, label, body, word_list, par, name):
    run = par.add_run("{}, Label: {}".format(count, str(label)))
    run.font.size = Pt(7)
    run.font.bold = True
    run.font.underline = True
    run.font.all_caps = True
    # run.font.outline = True
    if label >= 4:
        # run.font.highlight_color = WD_COLOR_INDEX.RED
        run.font.color.rgb = RGBColor.from_string('FF0000')
    elif label == 3:
        # run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.font.color.rgb = RGBColor.from_string('808000')
    elif label < 3:
        # run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        run.font.color.rgb = RGBColor.from_string('6897BB')
    run.add_break(WD_BREAK.LINE)
    body = ' '.join(body.split())
    if type(name) != type(' '):
        names = []
    else:
        names = name.split(";")
    # print (names)
    for word in word_list:
        f = word.split()[0]
        s = word.split()[1]
        if len(f) <= 3 and len(s) <= 3:
            word_list.remove(word)
    h_body = get_highlighted_body(body, word_list)
    h_body = unify_phrases(h_body)

    # print (h_body)
    r = re.compile('HSTART_')
    fragments = r.split(h_body)
    # par = doc.add_paragraph()
    second = fragments[0]
    done = False
    for n in names:
        if n in second:
            name_ind = second.lower().index(n.lower())
            run = par.add_run(second[:name_ind])
            run.font.size = Pt(6)
            run = par.add_run(second[name_ind:name_ind+len(n)])
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.underline = True
            run = par.add_run(second[name_ind+len(n):])
            run.font.size = Pt(6)
            done = True
            break
    if not done:
        run = par.add_run(fragments[0])
        run.font.size = Pt(6)
    

    # run.font.size = Pt(6)
    for ind, f in enumerate(fragments[1:]):
        first = f.split('HEND')[0]
        second = f.split('HEND')[1]
        # count = int(first[0])
        count = int(first[0])
        first = first[1:]

        run = par.add_run(first)
        font_map_obj = FONT_MAP[count]
        run.font.size = Pt(8)
        run.font.highlight_color = font_map_obj['highlight']
        run.font.color.rgb = font_map_obj['color']
        done =False
        if second :
            for n in names:
                if n in second:
                    name_ind = second.lower().index(n.lower())
                    run = par.add_run(second[:name_ind])
                    run.font.size = Pt(6)
                    run = par.add_run(second[name_ind:name_ind+len(n)])
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.underline = True
                    run = par.add_run(second[name_ind+len(n):])
                    run.font.size = Pt(6)
                    done = True
                    break
            if not done:
                run = par.add_run(second)
                run.font.size = Pt(6)

    return par



class1_count = 0
class2_count = 0
class3_count = 0
# try:
for ind, cl in enumerate(clusters):
    print (cl)
    try:
        if cl == -1:
            continue
        count = 0
        df_filt = df_res[df_res['cluster_label'] == cl]
        if len(df_filt) > 100:
            print (cl, len(df_filt))
            continue
        if avg_label_scores[ind] >= 3.5:
            class1_count += 1
            doc = doc1
            class1_clusters.append([cl,round(avg_label_scores[ind],2), round(cluster_metric[ind],2)])
        elif avg_label_scores[ind] < 3.5 and avg_label_scores[ind] >= 2.0:
            class2_count += 1
            doc = doc2
            class2_clusters.append([cl,round(avg_label_scores[ind],2), round(cluster_metric[ind],2)])
        else:
            class3_count += 1
            doc = doc3
            class3_clusters.append([cl,round(avg_label_scores[ind],2), round(cluster_metric[ind],2)])

        # par = doc.add_paragraph("Is it organized activity? (On a scale of 0-6, 6 being most organized) : ")
        # run = par.runs[0]
        # run.font.size = Pt(7)
        # par = doc.add_paragraph("Is it suspicious? (On a scale of 0-6, 6 being most trafficking) : ")
        # run = par.runs[0]
        # run.font.size = Pt(7)
        # par = doc.add_paragraph("Is it known or identified? (Y/N) : ")
        # run = par.runs[0]
        # run.font.size = Pt(7)
        # par = doc.add_paragraph("Additional comments.")
        # run = par.runs[0]
        # run.font.size = Pt(7)
        
        par = doc.add_paragraph('Cluster = {}'.format(cl))
        run = par.runs[0]
        run.font.size = Pt(11)
        par = doc.add_paragraph('{} ads, Average Label Score = {}, Cluster Metric = {}, {}'.format(len(df_filt),
            round(avg_label_scores[ind],2), round(cluster_metric[ind],2), top_bigrams[ind]))
        run = par.runs[0]
        run.font.size = Pt(7)
        
        
        table = doc.add_table(rows=1, cols=1)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Id, Label, Body'
        # hdr_cells[1].text = 'Original Label'
        # hdr_cells[2].text = 'Content'
        run = hdr_cells[0].paragraphs[0].runs[0]
        # hdr_cells[0].width = Inches(0.5)
        # hdr_cells[1].width = Inches(1)
        # hdr_cells[2].width = Inches(9)
        run.font.size = Pt(6)
        # run = hdr_cells[1].paragraphs[0].runs[0]
        # run.font.size = Pt(6)
        # run = hdr_cells[2].paragraphs[0].runs[0]
        # run.font.size = Pt(6)
        # hdr_cells[3].text = 'Title'
        # hdr_cells[4].text = 'Body'
        # print (ind, cl, len(df_filt))
        # print (cluster_word_list[ind])
        for i, row in df_filt.iterrows():
            count += 1
            row_cells = table.add_row().cells
            # row_cells[0].text = str(count)
            # row_cells[0].width = Inches(0.5)
            # row_cells[1].text = row['Name'] if type(row['Name']) == type(' ') else ' '
            # run = row_cells[0].paragraphs[0].runs[0]

            # run.font.size = Pt(6)
            # row_cells[1].text = str(row['label'])
            # run = row_cells[1].paragraphs[0].runs[0]
            # run.font.size = Pt(6)
            # row_cells[1].width = Inches(1)
            
            # row_cells[2].width = Inches(9)
            # row_cells[0].text = str(row['id'])
            # row_cells[1].text = row['Name'] if type(row['Name']) == type(' ') else ' '
            # row_cells[2].text = str(row['label'])
            # row_cells[3].text = row['title'] if type(row['title']) == type(' ') else ' '
            # row_cells[4].text = row['body'] if type(row['body']) == type(' ') else ' '
            # par = get_para_object(row['body'], cluster_word_list[ind].split(", "))

            par = row_cells[0].add_paragraph()
            get_para_object(str(count), row['label'], row['body'], cluster_word_list[ind].split(", ")[:20], par, row['Name'])
        doc.add_page_break()
    except Exception as e:
        print (e)
        pass
    # break
doc1.save(file_class_one)
doc2.save(file_class_two)
doc3.save(file_class_three)

with open(csv_class1, 'w') as f:
    writer = csv.writer(f)
    writer.writerows(class1_clusters)

with open(csv_class2, 'w') as f:
    writer = csv.writer(f)
    writer.writerows(class2_clusters)

with open(csv_class3, 'w') as f:
    writer = csv.writer(f)
    writer.writerows(class3_clusters)

print ("Class1 : {}, class2 : {}, class3: {}".format(class1_count, class2_count, class3_count))
